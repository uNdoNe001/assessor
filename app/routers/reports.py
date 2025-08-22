from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from sqlalchemy import select, func
from docx import Document
import os, uuid

from ..db import get_db
from .. import models
from ..deps import require_role

router = APIRouter(prefix="/reports", tags=["reports"])

OUT_DIR = "/app/generated"
os.makedirs(OUT_DIR, exist_ok=True)

@router.post("/generate")
def generate_report(client_id: int, assessment_id: int | None = None, db: Session = Depends(get_db), user=Depends(require_role(["pss_owner","pss_analyst","client_admin"]))):
    doc = Document()
    doc.add_heading("PSS Assessor Report", 0)
    doc.add_paragraph(f"Client ID: {client_id}")
    if assessment_id:
        doc.add_paragraph(f"Assessment ID: {assessment_id}")
        q = select(func.sum(models.Answer.maturity * models.Question.weight).label("wsum"), func.sum(models.Question.weight).label("w")).join(models.Question, models.Question.id==models.Answer.question_id).where(models.Answer.assessment_id==assessment_id)
        r = db.execute(q).first()
        avg = 0.0
        if r and r.w:
            avg = float(r.wsum or 0)/float(r.w or 1)
        doc.add_paragraph(f"Average Maturity (weighted): {avg:.2f}")
    fname = f"report_{client_id}_{assessment_id or 'na'}_{uuid.uuid4().hex[:8]}.docx"
    out_path = os.path.join(OUT_DIR, fname)
    doc.save(out_path)
    rec = models.Report(client_id=client_id, assessment_id=assessment_id, file_path=out_path)
    db.add(rec); db.commit()
    return {"file_path": out_path}
